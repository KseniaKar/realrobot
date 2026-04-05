"""
1. Для лота извлекаем attachmentId "Протокол аукциона" из страницы investmoscow.ru
2. Скачиваем .docx по API
3. Парсим: количество участников, победитель, цена
"""
import requests
import re
import json
import time
import os
import pandas as pd
from docx import Document

LOT_URL_TMPL = "https://investmoscow.ru/tenders/tender/{lot_id}"
DOC_URL_TMPL = "https://api.investmoscow.ru/investmoscow/tender/v1/tender/getattachedfilebyid?attachmentId={attach_id}"

headers_page = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
}

def get_protocol_attachment_id(lot_id):
    """Извлекаем attachmentId для 'Протокол аукциона' со страницы лота"""
    url = LOT_URL_TMPL.format(lot_id=lot_id)
    r = requests.get(url, headers=headers_page, timeout=15)
    
    # Ищем в JSON-данных страницы
    # Паттерн: "Протокол аукциона", ... attachmentId=XXXXXXX
    for m in re.finditer(r'Протокол\s*аукциона', r.text, re.IGNORECASE):
        # Ищем attachmentId в ближайших 500 символах
        start = max(0, m.start() - 500)
        context = r.text[start:m.end()]
        id_m = re.search(r'attachmentId[=: ]*(\d+)', context)
        if id_m:
            return int(id_m.group(1))
    
    return None

def parse_protocol_docx(filepath):
    """Парсим .docx протокол аукциона"""
    doc = Document(filepath)
    
    result = {
        "participants_count": 0,
        "winner": "",
        "winner_price": "",
        "second_place": "",
        "second_price": "",
    }
    
    all_text = "\n".join(p.text for p in doc.paragraphs)
    
    # Победитель
    m = re.search(r'Победителем.*?признан\s+(?:участник\s+)?(.+?),', all_text)
    if m:
        result["winner"] = m.group(1).strip().rstrip('"').rstrip("'")
    
    # Цена победителя
    m = re.search(r'наибольшую цену лота в размере\s+([\d\s,]+)', all_text)
    if m:
        result["winner_price"] = m.group(1).strip().rstrip(')')
    
    # Предпоследнее
    m = re.search(r'Участником, сделавшим предпоследнее.*?признан\s+(?:участник\s+)?(.+?),', all_text)
    if m:
        result["second_place"] = m.group(1).strip().rstrip('"').rstrip("'")
    
    m = re.search(r'предложивший цену лота в размере\s+([\d\s,]+)', all_text)
    if m:
        result["second_price"] = m.group(1).strip().rstrip(')')
    
    # Количество участников из таблицы
    for table in doc.tables:
        if len(table.rows) > 1 and len(table.columns) >= 2:
            header = table.rows[0].cells[0].text.strip()
            if "номер заявки" in header.lower() or "порядковый" in header.lower():
                result["participants_count"] = len(table.rows) - 1  # минус заголовок
                break
    
    return result

# --- ГЛАВНЫЙ КОД ---

# Загружаем CSV с лотами
df = pd.read_csv("data/investmoscow_completed_2026-04-04_geocoded.csv", encoding="utf-8-sig")

# Берём первые 5 лотов с platformLink (значит торги состоялись или были заявки)
test_lots = df[df["platformLink"].notna()].head(5)

print(f"Парсим {len(test_lots)} лотов...\n")

results = []

for idx, row in test_lots.iterrows():
    lot_id = int(row["номер_лота"])
    print(f"{'='*70}")
    print(f"Лот #{lot_id}")
    
    # 1. Находим attachmentId
    attach_id = get_protocol_attachment_id(lot_id)
    if not attach_id:
        print(f"  ⚠️  Не найден attachmentId для 'Протокол аукциона'")
        results.append({"номер_лота": lot_id, "error": "no_protocol"})
        continue
    
    print(f"  attachmentId: {attach_id}")
    
    # 2. Скачиваем .docx
    doc_url = DOC_URL_TMPL.format(attach_id=attach_id)
    filename = f"data/protocols/{lot_id}_protocol.docx"
    os.makedirs("data/protocols", exist_ok=True)
    
    try:
        r = requests.get(doc_url, headers=headers_page, timeout=15)
        if r.status_code != 200:
            print(f"  ⚠️  Ошибка скачивания: {r.status_code}")
            results.append({"номер_лота": lot_id, "error": f"download_{r.status_code}"})
            continue
        
        with open(filename, "wb") as f:
            f.write(r.content)
        print(f"  Скачано: {len(r.content)} bytes")
        
        # 3. Парсим
        parsed = parse_protocol_docx(filename)
        parsed["номер_лота"] = lot_id
        results.append(parsed)
        
        print(f"  Участников: {parsed['participants_count']}")
        print(f"  Победитель: {parsed['winner']}")
        print(f"  Цена: {parsed['winner_price']}")
        if parsed['second_place']:
            print(f"  2-е место: {parsed['second_place']} — {parsed['second_price']}")
        
    except Exception as e:
        print(f"  ❌ Ошибка: {e}")
        results.append({"номер_лота": lot_id, "error": str(e)})
    
    time.sleep(1)

# Итоговая таблица
print(f"\n\n{'='*70}")
print(f"ИТОГО")
print(f"{'='*70}")

for res in results:
    lot = res.get("номер_лота", "?")
    if "error" in res:
        print(f"  Лот #{lot}: ОШИБКА — {res['error']}")
    else:
        print(f"  Лот #{lot}: {res['participants_count']} участников, победитель: {res['winner'][:50]}, цена: {res['winner_price']}")
