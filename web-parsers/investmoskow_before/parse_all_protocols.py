"""
Восстановление кэша из уже скачанных .docx/.pdf файлов + продолжение парсинга
Парсим ТОЛЬКО лоты, где ещё нет данных об участниках
"""
import os
import re
import json
import time
import requests
import pandas as pd
from docx import Document
import pdfplumber
import io

PROTOCOLS_DIR = "data/protocols"
CACHE_FILE = os.path.join(PROTOCOLS_DIR, "protocol_cache.json")
OUTPUT_CSV = os.path.join(PROTOCOLS_DIR, "participants_data.csv")

LOT_URL_TMPL = "https://investmoscow.ru/tenders/tender/{lot_id}"
DOC_URL_TMPL = "https://api.investmoscow.ru/investmoscow/tender/v1/tender/getattachedfilebyid?attachmentId={aid}"

HEADERS = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
}

def parse_protocol_pdf(filepath):
    """Парсим .pdf протокол аукциона"""
    try:
        pdf = pdfplumber.open(filepath)
    except Exception as e:
        return {"error": f"Не удалось открыть pdf: {e}"}

    result = {
        "participants_count": 0,
        "winner": "",
        "winner_price": "",
        "winner_price_num": None,
        "second_place": "",
        "second_price": "",
        "second_price_num": None,
        "start_price": "",
        "auction_date": "",
        "auction_duration": "",
    }

    # Собираем весь текст
    all_text = "\n".join(page.extract_text() or "" for page in pdf.pages)

    # Начальная цена
    m = re.search(r'[Нн]ачальная цена.*?:\s*([\d\s,]+)', all_text)
    if m:
        result["start_price"] = m.group(1).strip()

    # Дата аукциона
    m_start = re.search(r'начала.*?(\d{2}\.\d{2}\.\d{4}\s+в\s+\d{2}:\d{2})', all_text)
    m_end = re.search(r'окончания.*?(\d{2}\.\d{2}\.\d{4}\s+в\s+\d{2}:\d{2})', all_text)
    if m_start and m_end:
        result["auction_duration"] = f"{m_start.group(1)} — {m_end.group(1)}"

    # Победитель
    m = re.search(r'[Пп]обедител.*?(?:признан|признается|определён)\s+(?:участник\s+)?(.+?)(?:\s+\(порядковый|,\s*предложивш)', all_text, re.DOTALL)
    if m:
        winner_raw = m.group(1).strip()
        # Убираем переносы строк и чистим
        winner_raw = re.sub(r'\s+', ' ', winner_raw)
        result["winner"] = winner_raw.rstrip('."').strip()

    # Цена победителя — ищем "наибольшую цену" или цену рядом с победителем
    m = re.search(r'наибольшую цену.*?([\d\s]+,\d{2})\s*руб', all_text, re.DOTALL)
    if not m:
        m = re.search(r'[Пп]обедител.*?цену.*?([\d\s]+,\d{2})\s*руб', all_text, re.DOTALL)
    if m:
        price_str = m.group(1).strip().rstrip(').').strip()
        result["winner_price"] = price_str
        price_num = re.sub(r'[^\d.]', '', price_str.replace(',', '.'))
        try:
            result["winner_price_num"] = float(price_num)
        except:
            pass

    # Предпоследнее место
    m = re.search(r'[Вв]торое место.*?(?:признан|признается|определён)\s+(?:участник\s+)?(.+?)(?:\s+\(порядковый|,\s*предложивш)', all_text, re.DOTALL)
    if m:
        second_raw = m.group(1).strip()
        second_raw = re.sub(r'\s+', ' ', second_raw)
        result["second_place"] = second_raw.rstrip('."').strip()

    m = re.search(r'[Вв]торое место.*?цену.*?([\d\s]+,\d{2})\s*руб', all_text, re.DOTALL)
    if m:
        price_str = m.group(1).strip().rstrip(').').strip()
        result["second_price"] = price_str
        price_num = re.sub(r'[^\d.]', '', price_str.replace(',', '.'))
        try:
            result["second_price_num"] = float(price_num)
        except:
            pass

    # Количество участников — считаем заявки в таблице или список
    # Для PDF таблица часто разбита на несколько страниц
    total_rows = 0
    header_found = False
    for page in pdf.pages:
        tables = page.extract_tables()
        for table in tables:
            if len(table) > 1 and len(table[0]) >= 1:
                header = str(table[0][0]).strip().lower()
                # Проверяем, что это таблица участников (первая строка — заголовок или номер)
                is_participants_table = (
                    "номер заявки" in header or
                    "порядковый" in header or
                    header == "№ заявки" or
                    header == "№" or
                    re.match(r'^\d{5,}$', str(table[0][0]).strip())  # Начинается с номера заявки
                )
                if is_participants_table:
                    header_found = True
                    # Если первая строка — заголовок, не считаем её
                    first_cell = str(table[0][0]).strip().lower()
                    start_idx = 1 if first_cell in ['№ заявки', '№', 'порядковый номер'] or 'заяв' in first_cell else 0
                    total_rows += len(table) - start_idx

    if header_found:
        result["participants_count"] = total_rows

    # Вариант 2: если таблицы не найдены, считаем строки с номерами заявок
    if result["participants_count"] == 0:
        participant_lines = re.findall(r'^\d{5,}\s+.+', all_text, re.MULTILINE)
        result["participants_count"] = len(participant_lines)

    pdf.close()
    return result


def parse_protocol_docx(filepath):
    """Парсим .docx протокол аукциона"""
    try:
        doc = Document(filepath)
    except Exception as e:
        return {"error": f"Не удалось открыть docx: {e}"}
    
    result = {
        "participants_count": 0,
        "winner": "",
        "winner_price": "",
        "winner_price_num": None,
        "second_place": "",
        "second_price": "",
        "second_price_num": None,
        "start_price": "",
        "auction_date": "",
        "auction_duration": "",
    }
    
    all_text = "\n".join(p.text for p in doc.paragraphs)
    
    # Начальная цена
    m = re.search(r'Начальная цена.*?:\s*([\d\s,]+)', all_text)
    if m:
        result["start_price"] = m.group(1).strip()
    
    # Дата аукциона
    m_start = re.search(r'начала.*?(\d{2}\.\d{2}\.\d{4}\s+в\s+\d{2}:\d{2})', all_text)
    m_end = re.search(r'окончания.*?(\d{2}\.\d{2}\.\d{4}\s+в\s+\d{2}:\d{2})', all_text)
    if m_start and m_end:
        result["auction_duration"] = f"{m_start.group(1)} — {m_end.group(1)}"
    
    # Победитель
    m = re.search(r'Победителем.*?признан\s+(?:участник\s+)?(.+?),\s*предложивш', all_text)
    if m:
        result["winner"] = m.group(1).strip().rstrip('."').strip()
    
    # Цена победителя
    m = re.search(r'наибольшую цену лота в размере\s+([\d\s,]+)', all_text)
    if m:
        price_str = m.group(1).strip().rstrip(').').strip()
        result["winner_price"] = price_str
        price_num = re.sub(r'[^\d.]', '', price_str.replace(',', '.'))
        try:
            result["winner_price_num"] = float(price_num)
        except:
            pass
    
    # Предпоследнее
    m = re.search(r'предпоследнее.*?признан\s+(?:участник\s+)?(.+?),\s*предложивш', all_text, re.DOTALL)
    if m:
        result["second_place"] = m.group(1).strip().rstrip('."').strip()
    
    m = re.search(r'предложивший цену лота в размере\s+([\d\s,]+)', all_text)
    if m:
        price_str = m.group(1).strip().rstrip(').').strip()
        result["second_price"] = price_str
        price_num = re.sub(r'[^\d.]', '', price_str.replace(',', '.'))
        try:
            result["second_price_num"] = float(price_num)
        except:
            pass
    
    # Количество участников из таблицы
    for table in doc.tables:
        if len(table.rows) > 1 and len(table.columns) >= 2:
            header = table.rows[0].cells[0].text.strip().lower()
            if "номер заявки" in header or "порядковый" in header or "№" in header:
                result["participants_count"] = len(table.rows) - 1
                break
    
    return result

def get_protocol_attachment_id(lot_id):
    """Извлекаем attachmentId для протокола со страницы лота.
    Приоритет: Протокол аукциона > Протокол рассмотрения заявок > любой протокол"""
    url = LOT_URL_TMPL.format(lot_id=lot_id)
    try:
        r = requests.get(url, headers=HEADERS, timeout=15)
        text = r.text

        # Ищем все attachmentId с названиями протоколов
        protocols = re.findall(r'\},\s*(\d{6,})\s*,\s*"([^"]*Протокол[^"]*)"', text)

        if not protocols:
            return None

        # Приоритет: аукциона > рассмотрения > любой
        for preferred in ['аукцион', 'рассмотрен']:
            for aid, name in protocols:
                if preferred in name.lower():
                    return int(aid)

        # Если ничего приоритетного не нашли — берём первый
        return int(protocols[0][0])
    except:
        pass
    return None

def main():
    print("=== ЗАГРУЗКА СУЩЕСТВУЮЩЕГО КЭША ===")
    
    cache = {}
    if os.path.exists(CACHE_FILE):
        try:
            with open(CACHE_FILE, "r", encoding="utf-8") as f:
                cache = json.load(f)
            print(f"Загружен кэш: {len(cache)} записей")
        except json.JSONDecodeError:
            print("Кэш повреждён, начинаем заново")
            cache = {}
    
    # Дополняем кэш из существующих файлов
    existing_files = [f for f in os.listdir(PROTOCOLS_DIR) if f.endswith(('.docx', '.pdf'))]
    new_files = 0
    for fname in existing_files:
        lot_id = fname.split('_')[0]
        if lot_id not in cache:
            try:
                filepath = os.path.join(PROTOCOLS_DIR, fname)
                if fname.endswith('.pdf'):
                    parsed = parse_protocol_pdf(filepath)
                else:
                    parsed = parse_protocol_docx(filepath)
                cache[lot_id] = parsed
                new_files += 1
            except Exception as e:
                cache[lot_id] = {"error": str(e)}
                new_files += 1

    if new_files > 0:
        print(f"Добавлено из файлов: {new_files}")

    print(f"Всего в кэше: {len(cache)}")

    # Сохраним кэш
    with open(CACHE_FILE, "w", encoding="utf-8") as f:
        json.dump(cache, f, ensure_ascii=False, indent=2)

    success_count = sum(1 for v in cache.values() if not v.get("error"))
    fail_count = sum(1 for v in cache.values() if v.get("error"))
    has_participants = sum(1 for v in cache.values() if not v.get("error") and v.get("participants_count", 0) > 0)
    no_participants = sum(1 for v in cache.values() if not v.get("error") and v.get("participants_count", 0) == 0)

    print(f"  Успешно: {success_count} (с участниками: {has_participants}, без: {no_participants})")
    print(f"  Ошибки: {fail_count}")

    # 2. Загружаем CSV и продолжаем парсинг
    print("\n=== ПРОДОЛЖЕНИЕ ПАРСИНГА (только лоты без участников) ===")
    df = pd.read_csv("data/investmoscow_completed_2026-04-04_geocoded.csv", encoding="utf-8-sig")
    lots = df[df["platformLink"].notna()]

    # Фильтруем: только те, которых НЕТ в кэше ИЛИ в кэше нет участников
    def need_parsing(lot_id):
        if lot_id not in cache:
            return True
        # Если в кэше есть, но participants_count == 0 или нет поля — перепарсим
        entry = cache[lot_id]
        # Не перепарсим лоты, где точно нет протокола (бессмысленно)
        if entry.get("error") == "no_protocol":
            return False
        if entry.get("error"):
            return True
        return entry.get("participants_count", 0) == 0

    lots_to_parse = [row for _, row in lots.iterrows() if need_parsing(str(int(row["номер_лота"])))]

    print(f"В кэше: {len(cache)} | Нужно распарсить: {len(lots_to_parse)}")

    for idx, row in enumerate(lots_to_parse, 1):
        lot_id = str(int(row["номер_лота"]))

        print(f"[{idx}/{len(lots_to_parse)}] Лот #{lot_id}", end=" ")

        # Находим attachmentId
        attach_id = get_protocol_attachment_id(lot_id)
        if not attach_id:
            print("❌ Нет протокола")
            cache[lot_id] = {"error": "no_protocol"}
            # Сохраняем каждые 10
            if idx % 10 == 0:
                with open(CACHE_FILE, "w", encoding="utf-8") as f:
                    json.dump(cache, f, ensure_ascii=False, indent=2)
            continue

        # Скачиваем
        doc_url = DOC_URL_TMPL.format(aid=attach_id)

        try:
            r = requests.get(doc_url, headers=HEADERS, timeout=15)
            if r.status_code != 200:
                print(f"❌ HTTP {r.status_code}")
                cache[lot_id] = {"error": f"http_{r.status_code}"}
                continue

            # Определяем тип файла по сигнатуре
            content_start = r.content[:8]
            if content_start[:4] == b'PK\x03\x04':
                ext = 'docx'
                filename = os.path.join(PROTOCOLS_DIR, f"{lot_id}_protocol.docx")
                parse_func = parse_protocol_docx
            elif content_start[:5] == b'%PDF-':
                ext = 'pdf'
                filename = os.path.join(PROTOCOLS_DIR, f"{lot_id}_protocol.pdf")
                parse_func = parse_protocol_pdf
            else:
                print(f"❌ Неизвестный формат (bytes: {content_start[:4]})")
                cache[lot_id] = {"error": "unknown_format"}
                continue

            with open(filename, "wb") as f:
                f.write(r.content)

            parsed = parse_func(filename)

            if parsed.get("error"):
                print(f"❌ {parsed['error']}")
            else:
                print(f"✅ [{ext}] {parsed['participants_count']} уч. — {parsed.get('winner','')[:40]}")

            cache[lot_id] = parsed

        except Exception as e:
            print(f"❌ {e}")
            cache[lot_id] = {"error": str(e)}
        
        # Сохраняем каждые 10
        if idx % 10 == 0:
            with open(CACHE_FILE, "w", encoding="utf-8") as f:
                json.dump(cache, f, ensure_ascii=False, indent=2)
        
        time.sleep(0.5)
    
    # Финальное сохранение
    with open(CACHE_FILE, "w", encoding="utf-8") as f:
        json.dump(cache, f, ensure_ascii=False, indent=2)
    
    # Статистика
    success_count = sum(1 for v in cache.values() if not v.get("error"))
    fail_count = sum(1 for v in cache.values() if v.get("error"))
    has_participants = sum(1 for v in cache.values() if not v.get("error") and v.get("participants_count", 0) > 0)
    no_participants = sum(1 for v in cache.values() if not v.get("error") and v.get("participants_count", 0) == 0)
    pdf_count = sum(1 for k in cache.keys() if os.path.exists(os.path.join(PROTOCOLS_DIR, f"{k}_protocol.pdf")))
    docx_count = sum(1 for k in cache.keys() if os.path.exists(os.path.join(PROTOCOLS_DIR, f"{k}_protocol.docx")))

    print(f"\n{'='*60}")
    print(f"ГОТОВО")
    print(f"  Всего: {len(cache)}")
    print(f"  Успешно: {success_count} (с участниками: {has_participants}, без: {no_participants})")
    print(f"  Ошибки: {fail_count}")
    print(f"  Файлов: PDF={pdf_count}, DOCX={docx_count}")

    # Сохраняем CSV
    results = []
    for lot_id, data in cache.items():
        results.append({"lot_id": lot_id, **data})
    res_df = pd.DataFrame(results)
    res_df.to_csv(OUTPUT_CSV, index=False, encoding="utf-8-sig")
    print(f"\n[FILE] {OUTPUT_CSV}")

    # Статистика по участникам
    ok = res_df[res_df.get("error").isna()] if "error" in res_df.columns else res_df
    if len(ok) > 0 and "participants_count" in ok.columns:
        counts = ok["participants_count"].dropna()
        if len(counts) > 0:
            print(f"\nУчастники: ср={counts.mean():.1f} мед={counts.median():.0f} макс={counts.max()} мин={counts.min()}")

if __name__ == "__main__":
    main()
