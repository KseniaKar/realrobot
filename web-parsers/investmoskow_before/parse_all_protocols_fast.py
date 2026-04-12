"""
Быстрый парсер протоколов investmoscow.ru — многопоточная версия
Парсим ТОЛЬКО лоты, где ещё нет данных об участниках
"""
import os
import re
import json
import time
import requests
import pandas as pd
from docx import Document
import pdfplumber
import io
from concurrent.futures import ThreadPoolExecutor, as_completed
from threading import Lock

PROTOCOLS_DIR = "data/protocols"
CACHE_FILE = os.path.join(PROTOCOLS_DIR, "protocol_cache.json")
OUTPUT_CSV = os.path.join(PROTOCOLS_DIR, "participants_data.csv")

LOT_URL_TMPL = "https://investmoscow.ru/tenders/tender/{lot_id}"
DOC_URL_TMPL = "https://api.investmoscow.ru/investmoscow/tender/v1/tender/getattachedfilebyid?attachmentId={aid}"

HEADERS = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
}

# Блокировка для потокобезопасной записи в кэш
cache_lock = Lock()

def parse_protocol_pdf(filepath):
    """Парсим .pdf протокол аукциона"""
    try:
        pdf = pdfplumber.open(filepath)
    except Exception as e:
        return {"error": f"Не удалось открыть pdf: {e}"}

    result = {
        "participants_count": 0,
        "winner": "",
        "winner_price": "",
        "winner_price_num": None,
        "second_place": "",
        "second_price": "",
        "second_price_num": None,
        "start_price": "",
        "auction_date": "",
        "auction_duration": "",
    }

    # Собираем весь текст
    all_text = "\n".join(page.extract_text() or "" for page in pdf.pages)

    # Начальная цена
    m = re.search(r'[Нн]ачальная цена.*?:\s*([\d\s,]+)', all_text)
    if m:
        result["start_price"] = m.group(1).strip()

    # Дата аукциона
    m_start = re.search(r'начала.*?(\d{2}\.\d{2}\.\d{4}\s+в\s+\d{2}:\d{2})', all_text)
    m_end = re.search(r'окончания.*?(\d{2}\.\d{2}\.\d{4}\s+в\s+\d{2}:\d{2})', all_text)
    if m_start and m_end:
        result["auction_duration"] = f"{m_start.group(1)} — {m_end.group(1)}"

    # Победитель
    m = re.search(r'[Пп]обедител.*?(?:признан|признается|определён)\s+(?:участник\s+)?(.+?)(?:\s+\(порядковый|,\s*предложивш)', all_text, re.DOTALL)
    if m:
        winner_raw = m.group(1).strip()
        winner_raw = re.sub(r'\s+', ' ', winner_raw)
        result["winner"] = winner_raw.rstrip('."').strip()

    # Цена победителя
    m = re.search(r'наибольшую цену.*?([\d\s]+,\d{2})\s*руб', all_text, re.DOTALL)
    if not m:
        m = re.search(r'[Пп]обедител.*?цену.*?([\d\s]+,\d{2})\s*руб', all_text, re.DOTALL)
    if m:
        price_str = m.group(1).strip().rstrip(').').strip()
        result["winner_price"] = price_str
        price_num = re.sub(r'[^\d.]', '', price_str.replace(',', '.'))
        try:
            result["winner_price_num"] = float(price_num)
        except:
            pass

    # Предпоследнее место
    m = re.search(r'[Вв]торое место.*?(?:признан|признается|определён)\s+(?:участник\s+)?(.+?)(?:\s+\(порядковый|,\s*предложивш)', all_text, re.DOTALL)
    if m:
        second_raw = m.group(1).strip()
        second_raw = re.sub(r'\s+', ' ', second_raw)
        result["second_place"] = second_raw.rstrip('."').strip()

    m = re.search(r'[Вв]торое место.*?цену.*?([\d\s]+,\d{2})\s*руб', all_text, re.DOTALL)
    if m:
        price_str = m.group(1).strip().rstrip(').').strip()
        result["second_price"] = price_str
        price_num = re.sub(r'[^\d.]', '', price_str.replace(',', '.'))
        try:
            result["second_price_num"] = float(price_num)
        except:
            pass

    # Количество участников — считаем заявки в таблице
    total_rows = 0
    header_found = False
    for page in pdf.pages:
        tables = page.extract_tables()
        for table in tables:
            if len(table) > 1 and len(table[0]) >= 1:
                header = str(table[0][0]).strip().lower()
                is_participants_table = (
                    "номер заявки" in header or
                    "порядковый" in header or
                    header == "№ заявки" or
                    header == "№" or
                    re.match(r'^\d{5,}$', str(table[0][0]).strip())
                )
                if is_participants_table:
                    header_found = True
                    first_cell = str(table[0][0]).strip().lower()
                    start_idx = 1 if first_cell in ['№ заявки', '№', 'порядковый номер'] or 'заяв' in first_cell else 0
                    total_rows += len(table) - start_idx

    if header_found:
        result["participants_count"] = total_rows

    # Вариант 2: если таблицы не найдены, считаем строки с номерами заявок
    if result["participants_count"] == 0:
        participant_lines = re.findall(r'^\d{5,}\s+.+', all_text, re.MULTILINE)
        result["participants_count"] = len(participant_lines)

    pdf.close()
    return result


def parse_protocol_docx(filepath):
    """Парсим .docx протокол аукциона"""
    try:
        doc = Document(filepath)
    except Exception as e:
        return {"error": f"Не удалось открыть docx: {e}"}

    result = {
        "participants_count": 0,
        "winner": "",
        "winner_price": "",
        "winner_price_num": None,
        "second_place": "",
        "second_price": "",
        "second_price_num": None,
        "start_price": "",
        "auction_date": "",
        "auction_duration": "",
    }

    all_text = "\n".join(p.text for p in doc.paragraphs)

    # Начальная цена
    m = re.search(r'Начальная цена.*?:\s*([\d\s,]+)', all_text)
    if m:
        result["start_price"] = m.group(1).strip()

    # Дата аукциона
    m_start = re.search(r'начала.*?(\d{2}\.\d{2}\.\d{4}\s+в\s+\d{2}:\d{2})', all_text)
    m_end = re.search(r'окончания.*?(\d{2}\.\d{2}\.\d{4}\s+в\s+\d{2}:\d{2})', all_text)
    if m_start and m_end:
        result["auction_duration"] = f"{m_start.group(1)} — {m_end.group(1)}"

    # Победитель
    m = re.search(r'Победителем.*?признан\s+(?:участник\s+)?(.+?),\s*предложивш', all_text)
    if m:
        result["winner"] = m.group(1).strip().rstrip('."').strip()

    # Цена победителя
    m = re.search(r'наибольшую цену лота в размере\s+([\d\s,]+)', all_text)
    if m:
        price_str = m.group(1).strip().rstrip(').').strip()
        result["winner_price"] = price_str
        price_num = re.sub(r'[^\d.]', '', price_str.replace(',', '.'))
        try:
            result["winner_price_num"] = float(price_num)
        except:
            pass

    # Предпоследнее
    m = re.search(r'предпоследнее.*?признан\s+(?:участник\s+)?(.+?),\s*предложивш', all_text, re.DOTALL)
    if m:
        result["second_place"] = m.group(1).strip().rstrip('."').strip()

    m = re.search(r'предложивший цену лота в размере\s+([\d\s,]+)', all_text)
    if m:
        price_str = m.group(1).strip().rstrip(').').strip()
        result["second_price"] = price_str
        price_num = re.sub(r'[^\d.]', '', price_str.replace(',', '.'))
        try:
            result["second_price_num"] = float(price_num)
        except:
            pass

    # Количество участников из таблицы
    table_found = False
    for table in doc.tables:
        if len(table.rows) > 1 and len(table.columns) >= 2:
            header_text = " ".join(c.text.strip().lower() for c in table.rows[0].cells)
            is_header = (
                "номер заявки" in header_text or
                "порядковый" in header_text or
                "номер" in header_text and "заявк" in header_text or
                "№ п/п" in header_text or
                "место" in header_text
            )
            if is_header:
                if len(table.rows) > 1:
                    first_data = table.rows[1].cells[0].text.strip()
                    if re.match(r'^\d+$', first_data) or re.match(r'^\d{5,}$', first_data):
                        result["participants_count"] = len(table.rows) - 1
                        table_found = True
                        break
            elif re.match(r'^\d{5,}$', table.rows[0].cells[0].text.strip()):
                result["participants_count"] = len(table.rows)
                table_found = True
                break

    # Если таблица не дала результат — проверяем текст
    if result["participants_count"] <= 1:
        applicant_mentions = re.findall(r'[Зз]аявка\s*№?\s*\d+', all_text)
        if applicant_mentions and len(applicant_mentions) > result["participants_count"]:
            result["participants_count"] = len(applicant_mentions)

    return result


def get_protocol_attachment_id(lot_id):
    """Извлекаем attachmentId для протокола со страницы лота."""
    url = LOT_URL_TMPL.format(lot_id=lot_id)
    try:
        r = requests.get(url, headers=HEADERS, timeout=15)
        text = r.text

        protocols = re.findall(r'\},\s*(\d{6,})\s*,\s*"([^"]*Протокол[^"]*)"', text)
        if not protocols:
            return None

        for preferred in ['аукцион', 'рассмотрен']:
            for aid, name in protocols:
                if preferred in name.lower():
                    return int(aid)

        return int(protocols[0][0])
    except:
        pass
    return None


def parse_single_lot(lot_id, cache):
    """Парсим один лот. Возвращает (lot_id, result, status_text)"""
    # Проверяем кэш
    if lot_id in cache:
        return lot_id, cache[lot_id], "SKIP"

    # Находим attachmentId
    attach_id = get_protocol_attachment_id(lot_id)
    if not attach_id:
        return lot_id, {"error": "no_protocol"}, "NO_PROTO"

    # Скачиваем
    doc_url = DOC_URL_TMPL.format(aid=attach_id)
    try:
        r = requests.get(doc_url, headers=HEADERS, timeout=15)
        if r.status_code != 200:
            return lot_id, {"error": f"http_{r.status_code}"}, f"HTTP_{r.status_code}"

        # Определяем тип файла по сигнатуре
        content_start = r.content[:8]
        if content_start[:4] == b'PK\x03\x04':
            ext = 'docx'
            filename = os.path.join(PROTOCOLS_DIR, f"{lot_id}_protocol.docx")
            parse_func = parse_protocol_docx
        elif content_start[:5] == b'%PDF-':
            ext = 'pdf'
            filename = os.path.join(PROTOCOLS_DIR, f"{lot_id}_protocol.pdf")
            parse_func = parse_protocol_pdf
        else:
            return lot_id, {"error": "unknown_format"}, "UNK_FMT"

        with open(filename, "wb") as f:
            f.write(r.content)

        parsed = parse_func(filename)
        if parsed.get("error"):
            return lot_id, parsed, f"ERR:{parsed['error'][:30]}"
        else:
            participants = parsed.get('participants_count', 0)
            winner = parsed.get('winner', '')[:30]
            return lot_id, parsed, f"OK:{ext}:{participants}уч"

    except Exception as e:
        return lot_id, {"error": str(e)[:100]}, f"EXC:{str(e)[:30]}"


def main():
    print("=== ЗАГРУЗКА СУЩЕСТВУЮЩЕГО КЭША ===")

    cache = {}
    if os.path.exists(CACHE_FILE):
        try:
            with open(CACHE_FILE, "r", encoding="utf-8") as f:
                cache = json.load(f)
            print(f"Загружен кэш: {len(cache)} записей")
        except json.JSONDecodeError:
            print("Кэш повреждён, начинаем заново")
            cache = {}

    # Дополняем кэш из существующих файлов
    existing_files = [f for f in os.listdir(PROTOCOLS_DIR) if f.endswith(('.docx', '.pdf'))]
    new_files = 0
    for fname in existing_files:
        lot_id = fname.split('_')[0]
        if lot_id not in cache:
            try:
                filepath = os.path.join(PROTOCOLS_DIR, fname)
                if fname.endswith('.pdf'):
                    parsed = parse_protocol_pdf(filepath)
                else:
                    parsed = parse_protocol_docx(filepath)
                cache[lot_id] = parsed
                new_files += 1
            except Exception as e:
                cache[lot_id] = {"error": str(e)}
                new_files += 1

    if new_files > 0:
        print(f"Добавлено из файлов: {new_files}")

    # Сохраняем кэш
    with open(CACHE_FILE, "w", encoding="utf-8") as f:
        json.dump(cache, f, ensure_ascii=False, indent=2)

    success_count = sum(1 for v in cache.values() if not v.get("error"))
    fail_count = sum(1 for v in cache.values() if v.get("error"))
    has_participants = sum(1 for v in cache.values() if not v.get("error") and v.get("participants_count", 0) > 0)
    no_participants = sum(1 for v in cache.values() if not v.get("error") and v.get("participants_count", 0) == 0)

    print(f"  Успешно: {success_count} (с участниками: {has_participants}, без: {no_participants})")
    print(f"  Ошибки: {fail_count}")

    # Загружаем CSV
    print("\n=== ПОДГОТОВКА К ПАРСИНГУ ===")
    csv_path = "data/investmoscow_completed_2022_2026_geocoded_mapped.csv"
    if not os.path.exists(csv_path):
        csv_path = "data/investmoscow_completed_2026-04-04_geocoded.csv"
    df = pd.read_csv(csv_path, encoding="utf-8-sig")
    lots = df[df["platformLink"].notna()]

    # Фильтруем: только те, которых НЕТ в кэше
    lots_to_parse = [str(int(row["номер_лота"])) for _, row in lots.iterrows() if str(int(row["номер_лота"])) not in cache]

    print(f"В кэше: {len(cache)} | Нужно распарсить: {len(lots_to_parse)}")
    print(f"Используем {min(10, len(lots_to_parse))} потоков")

    # Многопоточный парсинг
    processed = 0
    start_time = time.time()

    with ThreadPoolExecutor(max_workers=10) as executor:
        # Отправляем все задачи
        future_to_lot = {
            executor.submit(parse_single_lot, lot_id, cache): lot_id
            for lot_id in lots_to_parse
        }

        for future in as_completed(future_to_lot):
            lot_id, result, status = future.result()

            # Обновляем кэш потокобезопасно
            with cache_lock:
                cache[lot_id] = result

            processed += 1

            # Печатаем статус
            if status == "SKIP":
                print(f"[{processed}/{len(lots_to_parse)}] Лот #{lot_id} — пропущен (в кэше)")
            elif status == "NO_PROTO":
                print(f"[{processed}/{len(lots_to_parse)}] Лот #{lot_id} — ❌ нет протокола")
            else:
                print(f"[{processed}/{len(lots_to_parse)}] Лот #{lot_id} — ✅ {status}")

            # Сохраняем каждые 50 лотов
            if processed % 50 == 0:
                elapsed = time.time() - start_time
                speed = processed / elapsed if elapsed > 0 else 0
                remaining = len(lots_to_parse) - processed
                est_time = remaining / speed if speed > 0 else 0

                with open(CACHE_FILE, "w", encoding="utf-8") as f:
                    json.dump(cache, f, ensure_ascii=False, indent=2)

                print(f"\n>>> Прогресс: {processed}/{len(lots_to_parse)} | Скорость: {speed:.1f} лотов/с | Осталось: ~{est_time/60:.0f} мин\n")

    # Финальное сохранение
    with open(CACHE_FILE, "w", encoding="utf-8") as f:
        json.dump(cache, f, ensure_ascii=False, indent=2)

    # Статистика
    success_count = sum(1 for v in cache.values() if not v.get("error"))
    fail_count = sum(1 for v in cache.values() if v.get("error"))
    has_participants = sum(1 for v in cache.values() if not v.get("error") and v.get("participants_count", 0) > 0)
    no_participants = sum(1 for v in cache.values() if not v.get("error") and v.get("participants_count", 0) == 0)
    pdf_count = sum(1 for k in cache.keys() if os.path.exists(os.path.join(PROTOCOLS_DIR, f"{k}_protocol.pdf")))
    docx_count = sum(1 for k in cache.keys() if os.path.exists(os.path.join(PROTOCOLS_DIR, f"{k}_protocol.docx")))

    elapsed = time.time() - start_time

    print(f"\n{'='*60}")
    print(f"ГОТОВО за {elapsed/60:.1f} мин")
    print(f"  Всего: {len(cache)}")
    print(f"  Успешно: {success_count} (с участниками: {has_participants}, без: {no_participants})")
    print(f"  Ошибки: {fail_count}")
    print(f"  Файлов: PDF={pdf_count}, DOCX={docx_count}")

    # Сохраняем CSV
    results = []
    for lot_id, data in cache.items():
        results.append({"lot_id": lot_id, **data})
    res_df = pd.DataFrame(results)
    res_df.to_csv(OUTPUT_CSV, index=False, encoding="utf-8-sig")
    print(f"\n[FILE] {OUTPUT_CSV}")

    # Статистика по участникам
    ok = res_df[res_df.get("error").isna()] if "error" in res_df.columns else res_df
    if len(ok) > 0 and "participants_count" in ok.columns:
        counts = ok["participants_count"].dropna()
        if len(counts) > 0:
            print(f"\nУчастники: ср={counts.mean():.1f} мед={counts.median():.0f} макс={counts.max()} мин={counts.min()}")


if __name__ == "__main__":
    main()
